# applications/management/commands/import_surveys.py
from __future__ import annotations

import os
import re
import unicodedata
from dataclasses import dataclass
from typing import List, Tuple

from django.core.management.base import BaseCommand, CommandError
from django.db import transaction

from docx import Document

from applications.models import FormDefinition, Question, Choice


# --- CONFIG: map your 4 docs to slugs/names ---
SURVEYS = [
    {
        "slug": "SURVEY_E_PRIMER",
        "name": "Encuesta Inicial · Emprendedora",
        "docx_path": "Primer - E.docx",
    },
    {
        "slug": "SURVEY_E_FINAL",
        "name": "Encuesta Final · Emprendedora",
        "docx_path": "Final - E.docx",
    },
    {
        "slug": "SURVEY_M_PRIMER",
        "name": "Encuesta Inicial · Mentora",
        "docx_path": "Primer - M.docx",
    },
    {
        "slug": "SURVEY_M_FINAL",
        "name": "Encuesta Final · Mentora",
        "docx_path": "Final - M.docx",
    },
]


# --- Helpers ---
def _norm(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"\s+", " ", s)
    return s


def _strip_md(s: str) -> str:
    # your docs sometimes use * or ** for emphasis/required marks
    s = (s or "")
    s = s.replace("**", "")
    return s


def _is_required(qtext: str) -> bool:
    # treat an asterisk anywhere as required mark (common in forms)
    return "*" in (qtext or "")


def _clean_question_text(qtext: str) -> str:
    s = _strip_md(qtext)
    s = s.replace("*", "").strip()
    return _norm(s)


def _slugify(text: str) -> str:
    t = _clean_question_text(text).lower()
    t = unicodedata.normalize("NFD", t)
    t = "".join(ch for ch in t if unicodedata.category(ch) != "Mn")
    t = re.sub(r"[^a-z0-9]+", "_", t).strip("_")
    if not t:
        t = "q"
    return t[:60]


def _parse_options(opts_cell: str) -> List[str]:
    """
    Your docs typically store options separated by ';'
    Free-text is shown as '—'
    """
    s = _norm(opts_cell)
    if not s:
        return []
    if s == "—":
        return []
    # split by ';' if present, otherwise single option line
    parts = [p.strip() for p in s.split(";")]
    parts = [p for p in parts if p]
    return parts


def _infer_field_type(question_text: str, options: List[str]) -> str:
    """
    Very safe heuristic:
      - No options -> SHORT_TEXT (your docs use '—' for that)
      - Options exist:
          - "Sí/No" -> BOOLEAN
          - Many options or "seleccionar más de una" -> MULTI_CHOICE
          - Otherwise -> CHOICE
    """
    txt = _clean_question_text(question_text).lower()

    if not options:
        return Question.SHORT_TEXT

    opts_lower = [o.lower() for o in options]
    if len(options) == 2 and ("sí" in opts_lower[0] or "si" in opts_lower[0]) and ("no" in opts_lower[1]):
        return Question.BOOLEAN

    if "puedes seleccionar" in txt or "selecciona todas" in txt or "selecciona más de una" in txt:
        return Question.MULTI_CHOICE

    # Some questions are naturally multi even without that sentence; you can add more rules here.
    return Question.CHOICE


def _extract_docx_rows(docx_abs_path: str) -> List[Tuple[str, str]]:
    """
    Expects first table with header row, columns:
      - Question text
      - Options (or —)
    """
    doc = Document(docx_abs_path)
    if not doc.tables:
        raise CommandError(f"No tables found in DOCX: {docx_abs_path}")

    table = doc.tables[0]
    if len(table.rows) < 2:
        return []

    rows = []
    for row in table.rows[1:]:  # skip header
        q = _norm(row.cells[0].text)
        opts = _norm(row.cells[1].text)
        if not q:
            continue
        rows.append((q, opts))
    return rows


@transaction.atomic
def upsert_survey(form_slug: str, form_name: str, docx_abs_path: str) -> None:
    form_def, _ = FormDefinition.objects.update_or_create(
        slug=form_slug,
        defaults={"name": form_name},
    )

    # We will:
    #  - upsert questions by slug
    #  - overwrite choices to match the doc
    #  - set active=True and position in doc order
    rows = _extract_docx_rows(docx_abs_path)
    if not rows:
        raise CommandError(f"No question rows found in DOCX table: {docx_abs_path}")

    used_slugs = set()
    position = 1

    for q_raw, opts_raw in rows:
        required = _is_required(q_raw)
        q_text = _clean_question_text(q_raw)

        options = _parse_options(opts_raw)
        field_type = _infer_field_type(q_text, options)

        base_slug = _slugify(q_text)
        slug = base_slug
        i = 2
        while slug in used_slugs:
            slug = f"{base_slug}_{i}"
            i += 1
        used_slugs.add(slug)

        q_obj, _ = Question.objects.update_or_create(
            slug=slug,
            defaults={
                "text": q_text,
                "help_text": "",  # you can extend docs later to include help_text if needed
                "required": required,
                "active": True,
                "field_type": field_type,
                "position": position,
            },
        )

        # attach to this form (ManyToMany)
        form_def.questions.add(q_obj)

        # rewrite choices for choice / multi-choice
        if field_type in (Question.CHOICE, Question.MULTI_CHOICE, Question.BOOLEAN):
            # For BOOLEAN, we keep explicit choices as "yes/no" values, labels Sí/No.
            if field_type == Question.BOOLEAN:
                options_for_choice = ["Sí", "No"]
            else:
                options_for_choice = options

            # delete old choices then recreate in correct order
            Choice.objects.filter(question=q_obj).delete()
            pos2 = 1
            for opt in options_for_choice:
                label = _norm(opt)
                # generate a stable value
                val = unicodedata.normalize("NFD", label)
                val = "".join(ch for ch in val if unicodedata.category(ch) != "Mn")
                val = re.sub(r"[^a-zA-Z0-9]+", "_", val).strip("_").lower()
                if not val:
                    val = f"opt_{pos2}"

                Choice.objects.create(
                    question=q_obj,
                    label=label,
                    value=val,
                    position=pos2,
                )
                pos2 += 1

        position += 1

    # Optional: deactivate questions no longer in doc
    # We won’t do that automatically because you may have admin edits you want to keep.


class Command(BaseCommand):
    help = "Import/update the 4 Club Emprendo surveys from DOCX files."

    def add_arguments(self, parser):
        parser.add_argument(
            "--docs-dir",
            default=".",
            help="Directory where the DOCX files live (default: project root).",
        )

    def handle(self, *args, **options):
        docs_dir = options["docs_dir"]
        docs_dir = os.path.abspath(docs_dir)

        self.stdout.write(self.style.NOTICE(f"Using docs dir: {docs_dir}"))

        for s in SURVEYS:
            docx = os.path.join(docs_dir, s["docx_path"])
            if not os.path.exists(docx):
                raise CommandError(f"Missing DOCX: {docx}")

            self.stdout.write(self.style.NOTICE(f"Importing {s['slug']} from {s['docx_path']} ..."))
            upsert_survey(s["slug"], s["name"], docx)

        self.stdout.write(self.style.SUCCESS("✅ Surveys imported/updated successfully."))
